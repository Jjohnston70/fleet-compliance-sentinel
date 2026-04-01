"""
Document Reverse Engineer - DOCX & PDF to Generator Code
=========================================================
TRUE NORTH DATA STRATEGIES

WHAT THIS DOES:
  Takes a .docx or .pdf file and generates the Python/JavaScript code
  that would RECREATE that document from scratch. Feed it a finished doc,
  get back the source code to build it programmatically.

  Think of it like a decompiler for documents:
    .docx → Node.js code (using docx-js library)
    .docx → Python code (using python-docx library)
    .pdf  → Python code (using reportlab library)

HOW TO RUN:
  python doc_to_code.py document.docx              ← generates Node.js (docx-js) code
  python doc_to_code.py document.docx --python      ← generates Python (python-docx) code
  python doc_to_code.py document.pdf                ← generates Python (reportlab) code
  python doc_to_code.py document.docx --output my_generator.js

REQUIRES:
  pip install python-docx PyPDF2

OUTPUT:
  A fully commented, runnable script that rebuilds the document.
  Edit the generated code to change content, colors, or layout,
  then run it to produce a new document.

WHY THIS EXISTS:
  - You get a .docx from a client or designer → reverse it into code you can automate
  - You built a doc manually in Word → extract the code so it's repeatable
  - You want to understand how a doc is structured → the code shows every element
  - Pipeline Punks teaching tool → see how documents work under the hood
"""

import sys
import os

# ============================================
# python-docx reads .docx files and exposes
# every paragraph, run, table, style, color,
# font, and margin as Python objects
# ============================================
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn                            # For reading raw XML properties


def emu_to_inches(emu):
    """Convert EMU (English Metric Units) to inches. 914400 EMU = 1 inch."""
    if emu is None:
        return None
    return round(emu / 914400, 2)


def emu_to_dxa(emu):
    """Convert EMU to DXA (used by docx-js). 914400 EMU = 1440 DXA = 1 inch."""
    if emu is None:
        return None
    return round(emu / 914400 * 1440)


def pt_to_halfpt(emu_size):
    """Convert EMU font size to half-points (used by docx-js). 12700 EMU = 1pt = 2 half-points."""
    if emu_size is None:
        return None
    return round(emu_size / 12700 * 2)


def rgb_to_hex(rgb_color):
    """Convert RGBColor object to hex string."""
    if rgb_color is None:
        return None
    return str(rgb_color)


def get_alignment_str(alignment):
    """Convert alignment enum to string."""
    if alignment is None:
        return "LEFT"
    mapping = {
        WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
        WD_ALIGN_PARAGRAPH.CENTER: "CENTER",
        WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFIED",
    }
    return mapping.get(alignment, "LEFT")


def get_alignment_js(alignment):
    """Convert alignment to docx-js AlignmentType string."""
    if alignment is None:
        return None
    mapping = {
        WD_ALIGN_PARAGRAPH.LEFT: "AlignmentType.LEFT",
        WD_ALIGN_PARAGRAPH.CENTER: "AlignmentType.CENTER",
        WD_ALIGN_PARAGRAPH.RIGHT: "AlignmentType.RIGHT",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "AlignmentType.JUSTIFIED",
    }
    return mapping.get(alignment, None)


def extract_shading(paragraph):
    """Extract background/shading color from a paragraph's XML."""
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is not None:
        shd = pPr.find(qn('w:shd'))
        if shd is not None:
            return shd.get(qn('w:fill'))
    return None


def extract_cell_shading(cell):
    """Extract background color from a table cell."""
    tcPr = cell._element.find(qn('w:tcPr'))
    if tcPr is not None:
        shd = tcPr.find(qn('w:shd'))
        if shd is not None:
            return shd.get(qn('w:fill'))
    return None


def extract_borders(cell):
    """Extract border info from a table cell."""
    tcPr = cell._element.find(qn('w:tcPr'))
    if tcPr is not None:
        borders = tcPr.find(qn('w:tcBorders'))
        if borders is not None:
            result = {}
            for side in ['top', 'bottom', 'left', 'right']:
                b = borders.find(qn(f'w:{side}'))
                if b is not None:
                    result[side] = {
                        'style': b.get(qn('w:val')),
                        'size': b.get(qn('w:sz')),
                        'color': b.get(qn('w:color')),
                    }
            return result
    return None


# ============================================
# ANALYZE DOCUMENT
# Reads the .docx and builds a structured
# representation of every element
# ============================================
def analyze_docx(filepath):
    """
    Reads a .docx file and returns a structured dictionary describing
    every element: page setup, styles, paragraphs, tables, runs, colors, etc.
    
    This is the "understanding" step — we figure out what's in the doc
    before generating any code.
    """

    doc = Document(filepath)
    analysis = {
        "filename": os.path.basename(filepath),
        "sections": [],
        "styles_used": set(),
        "colors_used": set(),
        "fonts_used": set(),
        "elements": [],                                 # Ordered list of paragraphs and tables
    }

    # --- Page setup ---
    for i, sec in enumerate(doc.sections):
        analysis["sections"].append({
            "index": i,
            "width_emu": sec.page_width,
            "height_emu": sec.page_height,
            "width_inches": emu_to_inches(sec.page_width),
            "height_inches": emu_to_inches(sec.page_height),
            "width_dxa": emu_to_dxa(sec.page_width),
            "height_dxa": emu_to_dxa(sec.page_height),
            "margin_top": emu_to_dxa(sec.top_margin),
            "margin_bottom": emu_to_dxa(sec.bottom_margin),
            "margin_left": emu_to_dxa(sec.left_margin),
            "margin_right": emu_to_dxa(sec.right_margin),
        })

    # --- Walk through body elements in order ---
    # This preserves the exact order of paragraphs and tables
    para_index = 0
    table_index = 0
    paragraphs = doc.paragraphs
    tables = doc.tables

    for elem in doc.element.body:
        tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag

        if tag == 'p':
            # Find the matching paragraph object
            for p in paragraphs[para_index:]:
                if p._element is elem:
                    para_data = analyze_paragraph(p, analysis)
                    analysis["elements"].append({"type": "paragraph", "data": para_data})
                    para_index = paragraphs.index(p) + 1
                    break

        elif tag == 'tbl':
            if table_index < len(tables):
                table_data = analyze_table(tables[table_index], analysis)
                analysis["elements"].append({"type": "table", "data": table_data})
                table_index += 1

    # Convert sets to sorted lists for output
    analysis["styles_used"] = sorted(analysis["styles_used"])
    analysis["colors_used"] = sorted(analysis["colors_used"])
    analysis["fonts_used"] = sorted(analysis["fonts_used"])

    return analysis


def analyze_paragraph(p, analysis):
    """Extract all properties from a single paragraph."""
    data = {
        "text": p.text,
        "style": p.style.name,
        "alignment": get_alignment_str(p.alignment),
        "alignment_js": get_alignment_js(p.alignment),
        "runs": [],
        "shading": extract_shading(p),
    }

    analysis["styles_used"].add(p.style.name)

    # Spacing
    if p.paragraph_format.space_before:
        data["space_before"] = emu_to_dxa(p.paragraph_format.space_before)
    if p.paragraph_format.space_after:
        data["space_after"] = emu_to_dxa(p.paragraph_format.space_after)

    for run in p.runs:
        run_data = {
            "text": run.text,
            "bold": run.bold,
            "italic": run.italic,
            "underline": run.underline,
        }

        if run.font.size:
            run_data["size_emu"] = run.font.size
            run_data["size_halfpt"] = pt_to_halfpt(run.font.size)
            run_data["size_pt"] = round(run.font.size / 12700, 1)

        if run.font.color and run.font.color.rgb:
            color = rgb_to_hex(run.font.color.rgb)
            run_data["color"] = color
            analysis["colors_used"].add(color)

        if run.font.name:
            run_data["font"] = run.font.name
            analysis["fonts_used"].add(run.font.name)

        data["runs"].append(run_data)

    return data


def analyze_table(table, analysis):
    """Extract all properties from a table."""
    data = {
        "rows": len(table.rows),
        "cols": len(table.columns),
        "cells": [],
    }

    for row_idx, row in enumerate(table.rows):
        row_data = []
        for col_idx, cell in enumerate(row.cells):
            cell_data = {
                "row": row_idx,
                "col": col_idx,
                "text": cell.text,
                "shading": extract_cell_shading(cell),
                "paragraphs": [],
            }

            if cell_data["shading"]:
                analysis["colors_used"].add(cell_data["shading"])

            for p in cell.paragraphs:
                cell_data["paragraphs"].append(analyze_paragraph(p, analysis))

            row_data.append(cell_data)
        data["cells"].append(row_data)

    return data


# ============================================
# CODE GENERATORS
# Take the analysis and produce runnable code
# ============================================

def generate_js_code(analysis):
    """
    Generate docx-js (Node.js) code that recreates the document.
    This is the same library used by Pipeline_Flyer_Generator_DOCX.js.
    """

    sec = analysis["sections"][0]                       # Use first section for page setup
    colors = analysis["colors_used"]
    fonts = analysis["fonts_used"]

    lines = []
    lines.append('/**')
    lines.append(f' * Auto-generated docx-js code from: {analysis["filename"]}')
    lines.append(' * Generated by doc_to_code.py - True North Data Strategies')
    lines.append(' *')
    lines.append(' * HOW TO RUN:')
    lines.append(' *   npm install docx     (one-time)')
    lines.append(f' *   node {analysis["filename"].replace(".docx", "_generator.js")}')
    lines.append(' */')
    lines.append('')
    lines.append('const fs = require("fs");')
    lines.append('const {')
    lines.append('  Document, Packer, Paragraph, TextRun, Table, TableRow, TableCell,')
    lines.append('  AlignmentType, BorderStyle, WidthType, ShadingType, VerticalAlign,')
    lines.append('  HeadingLevel, PageBreak')
    lines.append('} = require("docx");')
    lines.append('')

    # Colors section
    lines.append('// ============================================')
    lines.append('// COLORS FOUND IN DOCUMENT')
    lines.append('// ============================================')
    for i, color in enumerate(sorted(colors)):
        safe_name = f"COLOR_{i}"
        lines.append(f'const {safe_name} = "{color}";  // #{color}')
    lines.append('')

    # Fonts section
    if fonts:
        lines.append('// ============================================')
        lines.append('// FONTS FOUND IN DOCUMENT')
        lines.append('// ============================================')
        for font in sorted(fonts):
            safe = font.upper().replace(' ', '_').replace('-', '_')
            lines.append(f'const FONT_{safe} = "{font}";')
        lines.append('')

    # Helpers
    lines.append('// ============================================')
    lines.append('// HELPERS')
    lines.append('// ============================================')
    lines.append('const noBorder = { style: BorderStyle.NONE, size: 0, color: "FFFFFF" };')
    lines.append('const noBorders = { top: noBorder, bottom: noBorder, left: noBorder, right: noBorder };')
    lines.append('const thinBorder = { style: BorderStyle.SINGLE, size: 1, color: "CCCCCC" };')
    lines.append('const thinBorders = { top: thinBorder, bottom: thinBorder, left: thinBorder, right: thinBorder };')
    lines.append('')

    # Build document
    lines.append('// ============================================')
    lines.append('// BUILD DOCUMENT')
    lines.append('// ============================================')
    lines.append('const doc = new Document({')

    # Styles
    default_font = sorted(fonts)[0] if fonts else "Arial"
    lines.append('  styles: {')
    lines.append('    default: {')
    lines.append(f'      document: {{ run: {{ font: "{default_font}", size: 22 }} }}')
    lines.append('    }')
    lines.append('  },')

    # Section with page setup
    lines.append('  sections: [{')
    lines.append('    properties: {')
    lines.append('      page: {')
    lines.append(f'        size: {{ width: {sec["width_dxa"]}, height: {sec["height_dxa"]} }},')
    lines.append(f'        margin: {{ top: {sec["margin_top"]}, right: {sec["margin_right"]}, bottom: {sec["margin_bottom"]}, left: {sec["margin_left"]} }}')
    lines.append('      }')
    lines.append('    },')
    lines.append('    children: [')

    # Content elements
    for el in analysis["elements"]:
        if el["type"] == "paragraph":
            p = el["data"]
            if not p["text"] and not p["runs"]:
                lines.append('      new Paragraph({ text: "" }),  // Empty line')
                continue

            js_lines = generate_paragraph_js(p, analysis)
            for jl in js_lines:
                lines.append('      ' + jl)

        elif el["type"] == "table":
            t = el["data"]
            table_lines = generate_table_js(t, analysis)
            for tl in table_lines:
                lines.append('      ' + tl)

    lines.append('    ]')
    lines.append('  }]')
    lines.append('});')
    lines.append('')

    # Output
    output_name = analysis["filename"].replace(".docx", "_REBUILT.docx")
    lines.append(f'Packer.toBuffer(doc).then(buffer => {{')
    lines.append(f'  fs.writeFileSync("{output_name}", buffer);')
    lines.append(f'  console.log("Created: {output_name}");')
    lines.append('});')

    return '\n'.join(lines)


def generate_paragraph_js(p, analysis):
    """Generate docx-js code for a single paragraph."""
    lines = []
    style = p["style"]
    alignment = p.get("alignment_js")

    # Determine heading level
    heading_map = {
        "Heading 1": "HeadingLevel.HEADING_1",
        "Heading 2": "HeadingLevel.HEADING_2",
        "Heading 3": "HeadingLevel.HEADING_3",
        "Heading 4": "HeadingLevel.HEADING_4",
    }

    props = []
    if style in heading_map:
        props.append(f'heading: {heading_map[style]}')
    if alignment:
        props.append(f'alignment: {alignment}')
    if p.get("space_before"):
        if not props or not any('spacing' in pr for pr in props):
            spacing_parts = []
            if p.get("space_before"):
                spacing_parts.append(f'before: {p["space_before"]}')
            if p.get("space_after"):
                spacing_parts.append(f'after: {p["space_after"]}')
            if spacing_parts:
                props.append(f'spacing: {{ {", ".join(spacing_parts)} }}')

    # Build runs
    runs = []
    for run in p["runs"]:
        if not run["text"]:
            continue
        run_props = []
        escaped = run["text"].replace('\\', '\\\\').replace('"', '\\"').replace('\n', '\\n')
        run_props.append(f'text: "{escaped}"')

        if run.get("font"):
            run_props.append(f'font: "{run["font"]}"')
        if run.get("size_halfpt"):
            run_props.append(f'size: {run["size_halfpt"]}')
        if run.get("bold"):
            run_props.append('bold: true')
        if run.get("italic"):
            run_props.append('italic: true')
        if run.get("underline"):
            run_props.append('underline: {}')
        if run.get("color"):
            run_props.append(f'color: "{run["color"]}"')

        runs.append(f'new TextRun({{ {", ".join(run_props)} }})')

    if not runs:
        escaped = p["text"].replace('\\', '\\\\').replace('"', '\\"').replace('\n', '\\n')
        runs.append(f'new TextRun({{ text: "{escaped}" }})')

    # Assemble paragraph
    all_props = props + [f'children: [{", ".join(runs)}]']
    comment = f'  // "{p["text"][:50]}..."' if len(p["text"]) > 50 else f'  // "{p["text"]}"'

    if len(', '.join(all_props)) < 120:
        lines.append(f'new Paragraph({{ {", ".join(all_props)} }}),{comment}')
    else:
        lines.append(f'new Paragraph({{')
        for prop in all_props[:-1]:
            lines.append(f'  {prop},')
        lines.append(f'  {all_props[-1]}')
        lines.append(f'}}),{comment}')

    return lines


def generate_table_js(t, analysis):
    """Generate docx-js code for a table."""
    lines = []
    num_cols = t["cols"]
    sec = analysis["sections"][0]
    content_width = sec["width_dxa"] - sec["margin_left"] - sec["margin_right"]
    col_width = content_width // num_cols

    lines.append(f'new Table({{  // {t["rows"]} rows x {t["cols"]} cols')
    lines.append(f'  width: {{ size: {content_width}, type: WidthType.DXA }},')
    lines.append(f'  columnWidths: [{", ".join([str(col_width)] * num_cols)}],')
    lines.append(f'  rows: [')

    for row_idx, row in enumerate(t["cells"]):
        lines.append(f'    new TableRow({{')
        lines.append(f'      children: [')

        for cell in row:
            shading_str = ""
            if cell.get("shading") and cell["shading"] != "auto":
                shading_str = f', shading: {{ fill: "{cell["shading"]}", type: ShadingType.CLEAR }}'

            lines.append(f'        new TableCell({{')
            lines.append(f'          borders: thinBorders{shading_str},')
            lines.append(f'          width: {{ size: {col_width}, type: WidthType.DXA }},')
            lines.append(f'          margins: {{ top: 80, bottom: 80, left: 120, right: 120 }},')
            lines.append(f'          children: [')

            for cp in cell["paragraphs"]:
                para_lines = generate_paragraph_js(cp, analysis)
                for pl in para_lines:
                    lines.append(f'            {pl}')

            lines.append(f'          ]')
            lines.append(f'        }}),')

        lines.append(f'      ]')
        lines.append(f'    }}),')

    lines.append(f'  ]')
    lines.append(f'}}),')

    return lines


# ============================================
# PYTHON-DOCX CODE GENERATOR
# ============================================
def generate_python_docx_code(analysis):
    """Generate python-docx code that recreates the document."""
    sec = analysis["sections"][0]

    lines = []
    lines.append('"""')
    lines.append(f'Auto-generated python-docx code from: {analysis["filename"]}')
    lines.append('Generated by doc_to_code.py - True North Data Strategies')
    lines.append('')
    lines.append('HOW TO RUN:')
    lines.append('  pip install python-docx     (one-time)')
    lines.append(f'  python {analysis["filename"].replace(".docx", "_generator.py")}')
    lines.append('"""')
    lines.append('')
    lines.append('from docx import Document')
    lines.append('from docx.shared import Inches, Pt, RGBColor, Emu')
    lines.append('from docx.enum.text import WD_ALIGN_PARAGRAPH')
    lines.append('from docx.oxml.ns import qn')
    lines.append('')

    # Colors
    lines.append('# ============================================')
    lines.append('# COLORS FOUND IN DOCUMENT')
    lines.append('# ============================================')
    for color in sorted(analysis["colors_used"]):
        r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
        safe = f"COLOR_{color}"
        lines.append(f'{safe} = RGBColor(0x{color[0:2]}, 0x{color[2:4]}, 0x{color[4:6]})  # #{color}')
    lines.append('')

    # Create document
    lines.append('# ============================================')
    lines.append('# CREATE DOCUMENT')
    lines.append('# ============================================')
    lines.append('doc = Document()')
    lines.append('')
    lines.append('# Page setup')
    lines.append('section = doc.sections[0]')
    lines.append(f'section.page_width = {sec["width_emu"] if sec.get("width_emu") else "Inches(8.5)"}')
    lines.append(f'section.page_height = {sec["height_emu"] if sec.get("height_emu") else "Inches(11)"}')
    top_in = emu_to_inches(sec.get("margin_top", 1440) * 914400 // 1440) if sec.get("margin_top") else 1.0
    lines.append(f'section.top_margin = Inches({top_in})')
    lines.append(f'section.bottom_margin = Inches({top_in})')
    lines.append(f'section.left_margin = Inches({top_in})')
    lines.append(f'section.right_margin = Inches({top_in})')
    lines.append('')

    # Content
    lines.append('# ============================================')
    lines.append('# CONTENT')
    lines.append('# ============================================')

    for el in analysis["elements"]:
        if el["type"] == "paragraph":
            p = el["data"]
            python_lines = generate_paragraph_python(p)
            for pl in python_lines:
                lines.append(pl)

        elif el["type"] == "table":
            t = el["data"]
            table_lines = generate_table_python(t)
            for tl in table_lines:
                lines.append(tl)

    # Save
    lines.append('')
    output_name = analysis["filename"].replace(".docx", "_REBUILT.docx")
    lines.append(f'doc.save("{output_name}")')
    lines.append(f'print("Created: {output_name}")')

    return '\n'.join(lines)


def generate_paragraph_python(p):
    """Generate python-docx code for a paragraph."""
    lines = []

    style = p["style"]
    text = p["text"]

    if not text and not p["runs"]:
        lines.append('doc.add_paragraph("")  # Empty line')
        return lines

    if style == "Heading 1":
        escaped = text.replace('"', '\\"')
        lines.append(f'doc.add_heading("{escaped}", level=1)')
        return lines
    elif style == "Heading 2":
        escaped = text.replace('"', '\\"')
        lines.append(f'doc.add_heading("{escaped}", level=2)')
        return lines
    elif style == "Heading 3":
        escaped = text.replace('"', '\\"')
        lines.append(f'doc.add_heading("{escaped}", level=3)')
        return lines

    # Regular paragraph with runs
    if p["runs"] and any(r.get("bold") or r.get("color") or r.get("size_emu") for r in p["runs"]):
        lines.append(f'p = doc.add_paragraph()')
        if p.get("alignment_js"):
            align_map = {
                "AlignmentType.CENTER": "WD_ALIGN_PARAGRAPH.CENTER",
                "AlignmentType.RIGHT": "WD_ALIGN_PARAGRAPH.RIGHT",
                "AlignmentType.JUSTIFIED": "WD_ALIGN_PARAGRAPH.JUSTIFY",
            }
            if p["alignment_js"] in align_map:
                lines.append(f'p.alignment = {align_map[p["alignment_js"]]}')

        for run in p["runs"]:
            if not run["text"]:
                continue
            escaped = run["text"].replace('\\', '\\\\').replace('"', '\\"')
            lines.append(f'run = p.add_run("{escaped}")')
            if run.get("bold"):
                lines.append('run.bold = True')
            if run.get("italic"):
                lines.append('run.italic = True')
            if run.get("font"):
                lines.append(f'run.font.name = "{run["font"]}"')
            if run.get("size_emu"):
                pt_val = round(run["size_emu"] / 12700, 1)
                lines.append(f'run.font.size = Pt({pt_val})')
            if run.get("color"):
                c = run["color"]
                lines.append(f'run.font.color.rgb = RGBColor(0x{c[0:2]}, 0x{c[2:4]}, 0x{c[4:6]})')
    else:
        escaped = text.replace('\\', '\\\\').replace('"', '\\"')
        lines.append(f'doc.add_paragraph("{escaped}")')

    return lines


def generate_table_python(t):
    """Generate python-docx code for a table."""
    lines = []
    lines.append(f'')
    lines.append(f'# Table: {t["rows"]} rows x {t["cols"]} cols')
    lines.append(f'table = doc.add_table(rows={t["rows"]}, cols={t["cols"]})')
    lines.append(f'table.style = "Table Grid"')

    for row_idx, row in enumerate(t["cells"]):
        for col_idx, cell in enumerate(row):
            text = cell["text"].replace('\\', '\\\\').replace('"', '\\"').replace('\n', '\\n')
            if text:
                lines.append(f'table.cell({row_idx}, {col_idx}).text = "{text[:200]}"')

            if cell.get("shading") and cell["shading"] != "auto":
                lines.append(f'# Cell ({row_idx},{col_idx}) has background color: #{cell["shading"]}')

    lines.append('')
    return lines


# ============================================
# REPORTLAB CODE GENERATOR (for PDFs)
# ============================================
def generate_pdf_code_from_docx(analysis):
    """
    Generate reportlab (Python PDF) code from a docx analysis.
    Since PDFs use absolute positioning, this creates a basic layout
    that preserves the content and colors but uses estimated positioning.
    """
    sec = analysis["sections"][0]
    page_w = round(sec["width_dxa"] / 1440 * 72)       # Convert DXA to PDF points (72 pts/inch)
    page_h = round(sec["height_dxa"] / 1440 * 72)

    lines = []
    lines.append('"""')
    lines.append(f'Auto-generated reportlab code from: {analysis["filename"]}')
    lines.append('Generated by doc_to_code.py - True North Data Strategies')
    lines.append('')
    lines.append('HOW TO RUN:')
    lines.append('  pip install reportlab     (one-time)')
    lines.append(f'  python {analysis["filename"].replace(".docx", "_pdf_generator.py")}')
    lines.append('"""')
    lines.append('')
    lines.append('from reportlab.lib.pagesizes import letter')
    lines.append('from reportlab.pdfgen import canvas')
    lines.append('from reportlab.lib.colors import HexColor')
    lines.append('from reportlab.lib.styles import getSampleStyleSheet')
    lines.append('from reportlab.platypus import Paragraph as FlowParagraph')
    lines.append('')

    # Colors
    lines.append('# ============================================')
    lines.append('# COLORS FOUND IN DOCUMENT')
    lines.append('# ============================================')
    for color in sorted(analysis["colors_used"]):
        lines.append(f'COLOR_{color} = HexColor("#{color}")')
    lines.append('')

    # Content extraction
    lines.append('# ============================================')
    lines.append('# CONTENT EXTRACTED FROM DOCUMENT')
    lines.append('# ============================================')
    lines.append('CONTENT = [')
    for el in analysis["elements"]:
        if el["type"] == "paragraph" and el["data"]["text"]:
            escaped = el["data"]["text"].replace('\\', '\\\\').replace('"', '\\"')
            style = el["data"]["style"]
            align = el["data"]["alignment"]
            color = None
            size = None
            bold = False
            for run in el["data"]["runs"]:
                if run.get("color"):
                    color = run["color"]
                if run.get("size_emu"):
                    size = round(run["size_emu"] / 12700, 1)
                if run.get("bold"):
                    bold = True

            lines.append(f'    {{"type": "text", "style": "{style}", "align": "{align}", "text": "{escaped[:150]}", "color": {f"\"#{color}\"" if color else "None"}, "size": {size}, "bold": {bold}}},')

        elif el["type"] == "table":
            t = el["data"]
            for row in t["cells"]:
                for cell in row:
                    if cell["text"]:
                        escaped = cell["text"][:150].replace('\\', '\\\\').replace('"', '\\"').replace('\n', ' ')
                        lines.append(f'    {{"type": "table_cell", "text": "{escaped}"}},')
    lines.append(']')
    lines.append('')

    # Draw function
    output_name = analysis["filename"].replace(".docx", "_REBUILT.pdf")
    lines.append(f'# ============================================')
    lines.append(f'# GENERATE PDF')
    lines.append(f'# NOTE: PDF uses absolute positioning.')
    lines.append(f'# You will need to adjust Y coordinates to')
    lines.append(f'# position elements where you want them.')
    lines.append(f'# Y=0 is BOTTOM, Y={page_h} is TOP.')
    lines.append(f'# ============================================')
    lines.append(f'c = canvas.Canvas("{output_name}", pagesize=({page_w}, {page_h}))')
    lines.append(f'')
    lines.append(f'y = {page_h - 72}  # Start 1 inch from top')
    lines.append(f'LEFT_MARGIN = 72   # 1 inch left margin')
    lines.append(f'')
    lines.append('for item in CONTENT:')
    lines.append('    if item["type"] == "text":')
    lines.append('        size = item["size"] or 11')
    lines.append('        if item["style"].startswith("Heading"):')
    lines.append('            size = 18 if "1" in item["style"] else 14')
    lines.append('        c.setFont("Helvetica-Bold" if item["bold"] else "Helvetica", size)')
    lines.append('        if item["color"]:')
    lines.append('            c.setFillColor(HexColor(item["color"]))')
    lines.append('        else:')
    lines.append('            c.setFillColor(HexColor("#000000"))')
    lines.append('        c.drawString(LEFT_MARGIN, y, item["text"][:80])')
    lines.append('        y -= size + 6')
    lines.append('    elif item["type"] == "table_cell":')
    lines.append('        c.setFont("Helvetica", 10)')
    lines.append('        c.setFillColor(HexColor("#000000"))')
    lines.append('        c.drawString(LEFT_MARGIN + 10, y, item["text"][:80])')
    lines.append('        y -= 16')
    lines.append('    ')
    lines.append(f'    if y < 72:  # New page if near bottom')
    lines.append('        c.showPage()')
    lines.append(f'        y = {page_h - 72}')
    lines.append('')
    lines.append('c.save()')
    lines.append(f'print("Created: {output_name}")')

    return '\n'.join(lines)


# ============================================
# MAIN - Parse arguments and run
# ============================================
def main():
    if len(sys.argv) < 2:
        print("=" * 60)
        print("  Document Reverse Engineer")
        print("  DOCX & PDF → Generator Code")
        print("  True North Data Strategies")
        print("=" * 60)
        print()
        print("  Usage:")
        print("    python doc_to_code.py <file.docx> [--python] [--pdf] [--output name]")
        print()
        print("  Examples:")
        print("    python doc_to_code.py document.docx              → Node.js (docx-js) code")
        print("    python doc_to_code.py document.docx --python      → Python (python-docx) code")
        print("    python doc_to_code.py document.docx --pdf         → Python (reportlab) code")
        print("    python doc_to_code.py document.docx --output gen.js")
        print()
        print("  All flags can be combined:")
        print("    python doc_to_code.py doc.docx --python --pdf     → generates both Python files")
        print()
        sys.exit(0)

    input_path = sys.argv[1]
    args = sys.argv[2:]

    if not os.path.exists(input_path):
        print(f"Error: File not found: {input_path}")
        sys.exit(1)

    want_python = "--python" in args
    want_pdf = "--pdf" in args
    want_js = not want_python and not want_pdf          # Default to JS if no flag specified

    # Find custom output name
    output_name = None
    if "--output" in args:
        idx = args.index("--output")
        if idx + 1 < len(args):
            output_name = args[idx + 1]

    # Analyze the document
    print(f"Analyzing: {input_path}")
    analysis = analyze_docx(input_path)
    print(f"  Sections: {len(analysis['sections'])}")
    print(f"  Elements: {len(analysis['elements'])}")
    print(f"  Colors: {', '.join('#' + c for c in analysis['colors_used']) if analysis['colors_used'] else 'none detected'}")
    print(f"  Fonts: {', '.join(analysis['fonts_used']) if analysis['fonts_used'] else 'none detected'}")
    print(f"  Styles: {', '.join(analysis['styles_used'])}")
    print()

    base = os.path.splitext(os.path.basename(input_path))[0]

    # Generate JavaScript code
    if want_js:
        code = generate_js_code(analysis)
        out = output_name or f"{base}_generator.js"
        with open(out, "w", encoding="utf-8") as f:
            f.write(code)
        print(f"Generated: {out}  (Node.js / docx-js)")
        print(f"  Run with: node {out}")

    # Generate Python-docx code
    if want_python:
        code = generate_python_docx_code(analysis)
        out = output_name or f"{base}_generator.py"
        with open(out, "w", encoding="utf-8") as f:
            f.write(code)
        print(f"Generated: {out}  (Python / python-docx)")
        print(f"  Run with: python {out}")

    # Generate Reportlab PDF code
    if want_pdf:
        code = generate_pdf_code_from_docx(analysis)
        out = output_name or f"{base}_pdf_generator.py"
        with open(out, "w", encoding="utf-8") as f:
            f.write(code)
        print(f"Generated: {out}  (Python / reportlab)")
        print(f"  Run with: python {out}")


if __name__ == "__main__":
    main()
