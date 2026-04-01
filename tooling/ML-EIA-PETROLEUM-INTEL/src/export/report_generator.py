from __future__ import annotations

import json
from pathlib import Path
from typing import Any

import pandas as pd


def _read_json(path: Path, fallback: Any) -> Any:
    if not path.exists():
        return fallback
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:  # noqa: BLE001
        return fallback


def _latest_forecasts(forecast_dir: Path) -> list[dict[str, Any]]:
    payloads_by_product: dict[str, dict[str, Any]] = {}
    latest_keys: dict[str, pd.Timestamp] = {}
    for p in sorted(forecast_dir.glob("*.json")):
        data = _read_json(p, {})
        product = data.get("product")
        if not product:
            continue

        generated_at = pd.to_datetime(data.get("generated_at"), errors="coerce", utc=True)
        if pd.isna(generated_at):
            generated_at = pd.Timestamp(p.stat().st_mtime, unit="s", tz="UTC")

        current = latest_keys.get(product)
        if current is not None and generated_at <= current:
            continue

        latest_keys[product] = generated_at
        payloads_by_product[product] = data
    return [payloads_by_product[k] for k in sorted(payloads_by_product.keys())]


def _make_forecast_chart(forecast_payload: dict[str, Any], chart_path: Path) -> None:
    import matplotlib.pyplot as plt  # lazy import

    points = forecast_payload.get("forecast", [])[:30]
    if not points:
        return
    x = [p["date"] for p in points]
    y = [p["point"] for p in points]

    plt.figure(figsize=(8, 3))
    plt.plot(x, y, linewidth=2)
    plt.xticks(rotation=45, ha="right")
    plt.title(f"30-Day Forecast: {forecast_payload.get('product', 'product')}")
    plt.tight_layout()
    chart_path.parent.mkdir(parents=True, exist_ok=True)
    plt.savefig(chart_path, dpi=140)
    plt.close()


def generate_executive_report(output_dir: Path, period: str = "monthly") -> Path:
    from docx import Document  # lazy import
    from docx.shared import Inches

    forecasts = _latest_forecasts(output_dir / "forecasts")
    analysis = _read_json(output_dir / "reports" / "analysis_snapshot.json", {})
    alerts = _read_json(output_dir / "alerts" / "active_alerts.json", {"alerts": []})

    doc = Document()
    doc.add_heading("Petroleum Executive Summary", 0)
    doc.add_paragraph(f"Period: {period}")
    doc.add_paragraph(f"Generated: {pd.Timestamp.utcnow().isoformat()}")

    regime = analysis.get("regime", {})
    doc.add_heading("Market Overview", level=1)
    doc.add_paragraph(
        f"Regime: {regime.get('current_regime', 'unknown')} | "
        f"Days in regime: {regime.get('days_in_regime', 'n/a')} | "
        f"Transition probability: {regime.get('transition_probability', 'n/a')}"
    )

    strategy = analysis.get("pricing_strategy", {})
    doc.add_heading("Pricing Strategy", level=1)
    doc.add_paragraph(f"Overall recommendation: {strategy.get('overall_recommendation', 'n/a')}")
    for product, info in strategy.get("products", {}).items():
        doc.add_paragraph(
            f"{product}: spread_pct={info.get('latest_spread_pct', 'n/a')} | "
            f"recommended={info.get('recommended_strategy', 'n/a')}"
        )

    doc.add_heading("Forecasts", level=1)
    for payload in forecasts:
        product = payload.get("product", "Unknown")
        metrics = payload.get("metrics", {})
        mape = metrics.get("mape")
        rmse = metrics.get("rmse")
        dir_acc = metrics.get("directional_accuracy")
        mape_str = f"{mape:.2f}" if isinstance(mape, (int, float)) else "n/a"
        rmse_str = f"{rmse:.4f}" if isinstance(rmse, (int, float)) else "n/a"
        dir_acc_str = f"{dir_acc:.2f}" if isinstance(dir_acc, (int, float)) else "n/a"
        doc.add_heading(product, level=2)
        doc.add_paragraph(f"MAPE={mape_str} | RMSE={rmse_str} | Directional Accuracy={dir_acc_str}")
        chart_path = output_dir / "reports" / "charts" / f"{product.lower().replace(' ', '_')}.png"
        try:
            _make_forecast_chart(payload, chart_path)
            if chart_path.exists():
                doc.add_picture(str(chart_path), width=Inches(6.5))
        except Exception:  # noqa: BLE001
            doc.add_paragraph("Chart generation skipped due to plotting error.")

    doc.add_heading("Active Alerts", level=1)
    for alert in alerts.get("alerts", [])[:12]:
        doc.add_paragraph(f"[{alert.get('severity', 'info')}] {alert.get('message', '')}")

    report_path = output_dir / "reports" / f"petroleum_executive_summary_{pd.Timestamp.utcnow().strftime('%Y%m%d')}.docx"
    report_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(report_path)
    return report_path
